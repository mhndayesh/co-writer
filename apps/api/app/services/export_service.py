"""Export a story to Markdown, plain text, DOCX, or a JSON bundle.

The JSON bundle shape matches Story Forge backup format so it round-trips.
"""
from __future__ import annotations

import io
from typing import Any

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.db.models import Chapter, Character, PlotThread, PlotThreadSceneLink, Revelation, SceneCard, Story, World


async def story_to_markdown(db: AsyncSession, story_id: str) -> str:
    story = await db.get(Story, story_id)
    world = await db.get(World, story_id)
    chapters = (await db.execute(select(Chapter).where(Chapter.story_id == story_id).order_by(Chapter.number))).scalars().all()
    characters = (await db.execute(select(Character).where(Character.story_id == story_id))).scalars().all()

    lines: list[str] = []
    lines.append(f"# {story.title or 'Untitled'}")
    if world and world.genre:
        lines.append(f"_{world.genre}_")
    lines.append("")
    if world and world.logline:
        lines.append(f"> {world.logline}")
        lines.append("")
    if characters:
        lines.append("## Cast")
        for c in characters:
            lines.append(f"- **{c.name}** — {c.role or 'unknown role'}")
        lines.append("")
    for ch in chapters:
        lines.append(f"## Chapter {ch.number}. {ch.title or 'Untitled'}")
        if ch.summary:
            lines.append(f"*{ch.summary}*")
            lines.append("")
        lines.append(ch.content or "")
        lines.append("")
    return "\n".join(lines)


async def story_to_bundle(db: AsyncSession, story_id: str) -> dict[str, Any]:
    """Story Forge-compatible JSON bundle."""
    from app.services import version_service

    v = await version_service.snapshot(db, story_id, note="export bundle")
    await db.commit()
    return {"app": "GInkNovelStudio", "version": 1, "story_id": story_id, "snapshot": v.snapshot}


async def story_to_docx_bytes(db: AsyncSession, story_id: str) -> bytes:
    from docx import Document

    md = await story_to_markdown(db, story_id)
    doc = Document()
    for line in md.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("> "):
            p = doc.add_paragraph(line[2:])
            p.style = "Intense Quote"
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("*") and line.endswith("*") and len(line) > 2:
            p = doc.add_paragraph()
            run = p.add_run(line.strip("*"))
            run.italic = True
        elif line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def import_bundle(db: AsyncSession, user_id: str, bundle: dict) -> str:
    """Import a Story Forge / G-Ink bundle and return the new story id."""
    snap = bundle.get("snapshot") or bundle
    world = snap.get("world", {}) or {}
    chars = snap.get("chars", []) or []
    chaps = snap.get("chaps", []) or []
    locs = snap.get("locations", []) or []
    facs = snap.get("factions", []) or []
    threads = snap.get("threads", []) or []
    scenes = snap.get("scenes", []) or []
    revelations = snap.get("revelations", []) or []
    thread_scene_links = snap.get("thread_scene_links", []) or []

    story = Story(
        user_id=user_id,
        title=world.get("title") or "Imported Story",
        genre=world.get("genre") or "",
    )
    db.add(story)
    await db.flush()

    db.add(World(
        story_id=story.id,
        title=world.get("title", ""),
        genre=world.get("genre", ""),
        logline=world.get("logline", ""),
        time_period=world.get("time_period", world.get("timePeriod", "")),
        setting=world.get("setting", ""),
        rules=list(world.get("rules") or []),
        themes=list(world.get("themes") or []),
        lore=world.get("lore", ""),
        seeds=world.get("seeds", ""),
    ))

    id_map: dict[str, str] = {}
    for c in chars:
        new_char = Character(
            story_id=story.id,
            name=c.get("name", "Unknown"),
            role=c.get("role", ""),
            icon=c.get("icon", ""),
            age=c.get("age", ""),
            appearance=c.get("appearance", ""),
            personality=c.get("personality", ""),
            backstory=c.get("backstory", ""),
            motivation=c.get("motivation", ""),
            flaw=c.get("flaw", ""),
            arc=c.get("arc", ""),
            status=c.get("status", "alive"),
        )
        db.add(new_char)
        await db.flush()
        if c.get("id"):
            id_map[c["id"]] = new_char.id

    from app.db.models import Location

    loc_id_map: dict[str, str] = {}
    for loc in locs:
        new_loc = Location(
            story_id=story.id,
            name=loc.get("name", ""),
            description=loc.get("description", ""),
            visual=loc.get("visual", ""),
        )
        db.add(new_loc)
        await db.flush()
        if loc.get("id"):
            loc_id_map[loc["id"]] = new_loc.id

    from app.db.models import Faction

    for f in facs:
        db.add(Faction(
            story_id=story.id,
            name=f.get("name", ""),
            description=f.get("description", ""),
            visual_signature=f.get("visual_signature", ""),
        ))

    chapter_id_map: dict[str, str] = {}
    for ch in chaps:
        new_char_ids = [id_map.get(cid, cid) for cid in (ch.get("characters") or []) if cid in id_map]
        loc = ch.get("location")
        new_loc = loc_id_map.get(loc) if loc else None
        pov = ch.get("pov")
        new_pov = id_map.get(pov) if pov else None
        new_chapter = Chapter(
            story_id=story.id,
            number=int(ch.get("number") or 1),
            title=ch.get("title", ""),
            content=ch.get("content", ""),
            summary=ch.get("summary", ""),
            pov_character_id=new_pov,
            location_id=new_loc,
            character_ids=new_char_ids,
            seeds=ch.get("seeds") or [],
        )
        db.add(new_chapter)
        await db.flush()
        if ch.get("id"):
            chapter_id_map[ch["id"]] = new_chapter.id

    thread_id_map: dict[str, str] = {}
    for t in threads:
        mapped_chapters = [chapter_id_map[cid] for cid in (t.get("chapter_ids") or []) if cid in chapter_id_map]
        new_thread = PlotThread(
            story_id=story.id,
            name=t.get("name", ""),
            status=t.get("status", "open"),
            description=t.get("description", ""),
            chapter_ids=mapped_chapters,
        )
        db.add(new_thread)
        await db.flush()
        if t.get("id"):
            thread_id_map[t["id"]] = new_thread.id

    scene_id_map: dict[str, str] = {}
    for s in scenes:
        old_threads = s.get("plot_threads") or s.get("plot_thread_ids") or []
        old_chars = s.get("characters") or s.get("character_ids") or []
        new_scene = SceneCard(
            story_id=story.id,
            chapter_id=chapter_id_map.get(s.get("chapter_id") or ""),
            ordinal=int(s.get("ordinal") or 0),
            beat=s.get("beat", ""),
            title=s.get("title", ""),
            summary=s.get("summary", ""),
            goal=s.get("goal", ""),
            conflict=s.get("conflict", ""),
            outcome=s.get("outcome", ""),
            pov_character_id=id_map.get(s.get("pov") or s.get("pov_character_id") or ""),
            location_id=loc_id_map.get(s.get("location") or s.get("location_id") or ""),
            character_ids=[id_map[cid] for cid in old_chars if cid in id_map],
            plot_thread_ids=[thread_id_map[tid] for tid in old_threads if tid in thread_id_map],
            time_anchor=s.get("time_anchor", ""),
            time_sort_key=s.get("time_sort_key"),
            duration_hint=s.get("duration_hint", ""),
            sensory_palette=s.get("sensory_palette") or {},
            source_excerpt=s.get("source_excerpt", ""),
            content=s.get("content", ""),
        )
        db.add(new_scene)
        await db.flush()
        if s.get("id"):
            scene_id_map[s["id"]] = new_scene.id

    for r in revelations:
        db.add(Revelation(
            story_id=story.id,
            scene_id=scene_id_map.get(r.get("scene_id") or ""),
            chapter_id=chapter_id_map.get(r.get("chapter_id") or ""),
            description=r.get("description", ""),
            kind=r.get("kind", "revelation"),
            characters_who_know=[id_map[cid] for cid in (r.get("characters_who_know") or []) if cid in id_map],
            reader_knows=bool(r.get("reader_knows", False)),
            notes=r.get("notes", ""),
            confidence=float(r.get("confidence", 1.0) or 1.0),
        ))

    for link in thread_scene_links:
        thread_id = thread_id_map.get(link.get("thread_id") or "")
        scene_id = scene_id_map.get(link.get("scene_id") or "")
        if not thread_id or not scene_id:
            continue
        db.add(PlotThreadSceneLink(
            story_id=story.id,
            thread_id=thread_id,
            scene_id=scene_id,
            chapter_id=chapter_id_map.get(link.get("chapter_id") or ""),
            status=link.get("status", "touch"),
            strength=float(link.get("strength", 1.0) or 1.0),
            evidence=link.get("evidence", ""),
        ))

    await db.commit()
    return story.id
