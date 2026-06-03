# apps/api/app/services/export_service.py
#
# Generates downloadable files from story data.
# Dependencies (add to pyproject.toml):
#   reportlab>=4.0, ebooklib>=0.18, python-docx>=1.1, Pillow>=10.0
#
# Each function returns raw bytes + (filename, mimetype) tuple.

from __future__ import annotations

import html as _html
import io
import re
import zipfile
from datetime import datetime
from typing import Optional

# ── PDF ─────────────────────────────────────────────────────────────────────
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    BaseDocTemplate, Paragraph, Spacer, PageBreak,
    Frame, PageTemplate,
)
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

# ── EPUB ─────────────────────────────────────────────────────────────────────
from ebooklib import epub

# ── DOCX ─────────────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _word_count(text: str) -> int:
    return len(re.findall(r"\w+", text))


def _clean_text(text: str) -> str:
    """Strip markdown-style formatting for plain-text exports."""
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*",     r"\1", text)
    text = re.sub(r"#+ ",           "",    text)
    return text.strip()


def _pdf_text(text: str) -> str:
    """Clean + XML-escape author prose for ReportLab `Paragraph`.

    Paragraph parses an HTML-ish markup mini-language, so a literal `<` or `&` in
    the prose (e.g. "5 < 10", "Tom & Jerry") raises an unhandled error mid-export
    (and `<i>`-style fragments would inject formatting). Escaping `&<>` makes it
    inert — matching what the EPUB path already does."""
    return _html.escape(_clean_text(text), quote=False)


def _paragraphs(content: str) -> list[str]:
    """Split chapter content into paragraphs, dropping empties."""
    return [p.strip() for p in content.split("\n\n") if p.strip()]


# ---------------------------------------------------------------------------
# PDF Export
# ---------------------------------------------------------------------------

def export_pdf(
    title: str,
    author: str,
    tagline: Optional[str],
    chapters: list[dict],  # [{"number": int, "title": str, "content": str}]
) -> tuple[bytes, str, str]:
    """Returns (bytes, filename, mimetype)."""
    buf = io.BytesIO()

    doc = BaseDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=3.5 * cm,
        rightMargin=3 * cm,
        topMargin=3 * cm,
        bottomMargin=3 * cm,
    )

    frame = Frame(
        doc.leftMargin, doc.bottomMargin,
        doc.width, doc.height, id="main",
    )

    def _header_footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Times-Italic", 9)
        canvas.setFillColor(colors.HexColor("#888888"))
        if doc.page > 1:
            canvas.drawString(doc.leftMargin, doc.bottomMargin - 1 * cm,
                              f"{title} — {author}")
            canvas.drawRightString(doc.width + doc.leftMargin,
                                   doc.bottomMargin - 1 * cm,
                                   str(doc.page))
        canvas.restoreState()

    template = PageTemplate(id="main", frames=frame, onPage=_header_footer)
    doc.addPageTemplates([template])

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "BookTitle",
        parent=styles["Heading1"],
        fontSize=28,
        leading=36,
        alignment=TA_CENTER,
        spaceAfter=12,
        fontName="Times-Bold",
    )
    author_style = ParagraphStyle(
        "AuthorLine",
        parent=styles["Normal"],
        fontSize=14,
        alignment=TA_CENTER,
        spaceAfter=8,
        textColor=colors.HexColor("#555555"),
        fontName="Times-Italic",
    )
    chapter_heading = ParagraphStyle(
        "ChapterHeading",
        parent=styles["Heading2"],
        fontSize=16,
        leading=22,
        spaceBefore=24,
        spaceAfter=16,
        fontName="Times-Bold",
        alignment=TA_CENTER,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=11,
        leading=17,
        firstLineIndent=24,
        alignment=TA_JUSTIFY,
        fontName="Times-Roman",
        spaceAfter=6,
    )
    tagline_style = ParagraphStyle(
        "Tagline",
        parent=styles["Normal"],
        fontSize=12,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#333333"),
        fontName="Times-Italic",
        spaceAfter=48,
    )

    story: list = []

    # Cover page
    story.append(Spacer(1, 6 * cm))
    story.append(Paragraph(_pdf_text(title), title_style))
    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph(_html.escape(author, quote=False), author_style))
    if tagline:
        story.append(Spacer(1, 1 * cm))
        story.append(Paragraph(f"“{_pdf_text(tagline)}”", tagline_style))
    story.append(PageBreak())

    # Chapters
    for ch in chapters:
        story.append(Paragraph(
            f"Chapter {ch['number']}: {_pdf_text(ch['title'])}",
            chapter_heading,
        ))
        for para in _paragraphs(ch["content"]):
            story.append(Paragraph(_pdf_text(para), body_style))
        story.append(PageBreak())

    doc.build(story)
    filename = re.sub(r"[^\w\-]", "_", title.lower())
    return buf.getvalue(), f"{filename}.pdf", "application/pdf"


# ---------------------------------------------------------------------------
# EPUB Export
# ---------------------------------------------------------------------------

def export_epub(
    title: str,
    author: str,
    tagline: Optional[str],
    genre: Optional[str],
    chapters: list[dict],
) -> tuple[bytes, str, str]:
    book = epub.EpubBook()
    book.set_identifier(f"ginink-{re.sub(r'[^a-z0-9]', '-', title.lower())}")
    book.set_title(title)
    book.set_language("en")
    book.add_author(author)
    if genre:
        book.add_metadata("DC", "subject", genre)

    css = """
        body { font-family: Georgia, serif; font-size: 1em; line-height: 1.7;
               margin: 2em; color: #1a1a1a; }
        h1 { text-align: center; font-size: 1.8em; margin-bottom: 0.3em; }
        h2 { text-align: center; font-size: 1.3em; margin-top: 2em; margin-bottom: 1em;
             border-bottom: 1px solid #ccc; padding-bottom: 0.4em; }
        p  { text-indent: 1.5em; margin: 0; }
        p.first { text-indent: 0; }
        .tagline { text-align: center; font-style: italic; color: #555;
                   margin: 2em 3em; font-size: 1.05em; }
    """
    style_item = epub.EpubItem(
        uid="style", file_name="style/main.css",
        media_type="text/css", content=css,
    )
    book.add_item(style_item)

    # Escape all user-supplied strings before embedding in XHTML.
    # Without this a story title like <script>...</script> would execute in
    # e-readers or any browser-based EPUB preview (stored XSS).
    esc_title = _html.escape(title)
    esc_author = _html.escape(author)
    esc_tagline = _html.escape(tagline) if tagline else ""

    # Cover page
    cover_html = (
        f"<h1>{esc_title}</h1><p style='text-align:center'>by {esc_author}</p>"
        + (f"<p class='tagline'>{esc_tagline}</p>" if esc_tagline else "")
    )
    cover_ch = epub.EpubHtml(title="Cover", file_name="cover.xhtml", lang="en")
    cover_ch.content = f"<html><body>{cover_html}</body></html>"
    cover_ch.add_item(style_item)
    book.add_item(cover_ch)

    epub_chapters = [cover_ch]
    toc = []

    for ch in chapters:
        paras = _paragraphs(ch["content"])
        para_html = ""
        for i, p in enumerate(paras):
            cls = "first" if i == 0 else ""
            para_html += f'<p class="{cls}">{_html.escape(_clean_text(p))}</p>\n'

        ch_title_raw = f"Chapter {ch['number']}: {_clean_text(ch['title'])}"
        ch_title_esc = _html.escape(ch_title_raw)
        epub_ch = epub.EpubHtml(
            title=ch_title_raw,  # ebooklib metadata: plain text, not rendered as HTML
            file_name=f"chapter_{ch['number']:03d}.xhtml",
            lang="en",
        )
        epub_ch.content = (
            f"<html><body>"
            f"<h2>{ch_title_esc}</h2>"
            f"{para_html}"
            f"</body></html>"
        )
        epub_ch.add_item(style_item)
        book.add_item(epub_ch)
        epub_chapters.append(epub_ch)
        toc.append(epub.Link(epub_ch.file_name, ch_title_raw, f"ch{ch['number']}"))

    book.toc = toc
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav"] + epub_chapters

    buf = io.BytesIO()
    epub.write_epub(buf, book, {})
    filename = re.sub(r"[^\w\-]", "_", title.lower())
    return buf.getvalue(), f"{filename}.epub", "application/epub+zip"


# ---------------------------------------------------------------------------
# DOCX Export (Shunn Standard Manuscript Format)
# ---------------------------------------------------------------------------

def export_docx(
    title: str,
    author: str,
    chapters: list[dict],
) -> tuple[bytes, str, str]:
    doc = Document()

    # Page setup: 1-inch margins, double spacing
    section = doc.sections[0]
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)

    FONT_NAME = "Courier New"
    FONT_SIZE = Pt(12)

    def _set_font(run):
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE

    def _para(text: str, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False) -> None:
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.line_spacing = Pt(24)   # double-spaced at 12pt
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(text)
        _set_font(run)
        run.bold = bold
        return p

    total_words = sum(_word_count(ch["content"]) for ch in chapters)
    word_est = f"~{round(total_words / 1000)}k words" if total_words > 999 else f"{total_words} words"

    # Title page header (top right)
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.paragraph_format.space_after = Pt(0)
    run = header_p.add_run(f"{author}\n{word_est}")
    _set_font(run)

    doc.add_paragraph()  # spacer

    # Title block (centered, ~halfway down page)
    for _ in range(8):
        _para("")

    _para(title.upper(), align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    _para("by", align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(author, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    for ch in chapters:
        ch_title = f"Chapter {ch['number']}: {_clean_text(ch['title'])}"
        # Chapter heading centered, 1/3 down page
        for _ in range(10):
            _para("")
        _para(ch_title.upper(), align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        doc.add_paragraph()

        first = True
        for para_text in _paragraphs(ch["content"]):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = Pt(24)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.first_line_indent = Inches(0) if first else Inches(0.5)
            run = p.add_run(_clean_text(para_text))
            _set_font(run)
            first = False

        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    filename = re.sub(r"[^\w\-]", "_", title.lower())
    return buf.getvalue(), f"{filename}_manuscript.docx", \
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


# ---------------------------------------------------------------------------
# Submission Package (zip: PDF + EPUB + DOCX + synopsis TXT)
# ---------------------------------------------------------------------------

async def export_submission_package(
    title: str,
    author: str,
    tagline: Optional[str],
    genre: Optional[str],
    chapters: list[dict],
    synopsis: Optional[str] = None,
) -> tuple[bytes, str, str]:
    pdf_bytes,  pdf_name,  _ = export_pdf(title, author, tagline, chapters)
    epub_bytes, epub_name, _ = export_epub(title, author, tagline, genre, chapters)
    docx_bytes, docx_name, _ = export_docx(title, author, chapters)

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(pdf_name,  pdf_bytes)
        zf.writestr(epub_name, epub_bytes)
        zf.writestr(docx_name, docx_bytes)

        if synopsis:
            zf.writestr("synopsis.txt", synopsis.encode("utf-8"))

        # README
        readme = (
            f"Submission Package — {title}\n"
            f"Author: {author}\n"
            f"Genre: {genre or 'unspecified'}\n"
            f"Generated: {datetime.utcnow().strftime('%Y-%m-%d')}\n\n"
            f"Files included:\n"
            f"  {pdf_name}  — reading copy\n"
            f"  {epub_name} — e-reader format\n"
            f"  {docx_name} — Shunn standard manuscript\n"
            + ("  synopsis.txt     — one-page synopsis\n" if synopsis else "")
        )
        zf.writestr("README.txt", readme.encode("utf-8"))

    filename = re.sub(r"[^\w\-]", "_", title.lower())
    return buf.getvalue(), f"{filename}_submission.zip", "application/zip"
